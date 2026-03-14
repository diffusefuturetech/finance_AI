"""Export stock analysis reports to Word (.docx) format."""

import logging
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

logger = logging.getLogger(__name__)

# Style constants
HEADER_BG = RGBColor(0x1A, 0x56, 0xDB)
ALT_ROW_BG = RGBColor(0xF0, 0xF4, 0xFF)
COLOR_RED = RGBColor(0xEF, 0x53, 0x50)
COLOR_GREEN = RGBColor(0x26, 0xA6, 0x9A)
COLOR_ORANGE = RGBColor(0xFF, 0x98, 0x00)
COLOR_BLUE = RGBColor(0x1A, 0x56, 0xDB)
COLOR_GRAY = RGBColor(0x99, 0x99, 0x99)


class DocxExporter:
    """Generate Word documents from stock analysis data."""

    def __init__(self, output_dir: Path | None = None):
        self.output_dir = output_dir or Path.cwd()
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def _set_cell_shading(self, cell, color: RGBColor):
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), f"{color}")
        shading.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shading)

    def _add_styled_table(self, doc, headers, rows):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(10)
            self._set_cell_shading(cell, HEADER_BG)
        for row_idx, row_data in enumerate(rows):
            for col_idx, value in enumerate(row_data):
                cell = table.rows[row_idx + 1].cells[col_idx]
                cell.text = str(value)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.size = Pt(9)
                if row_idx % 2 == 1:
                    self._set_cell_shading(cell, ALT_ROW_BG)
        return table

    def _add_colored_table(self, doc, headers, rows, color_rules=None):
        """Table with per-cell conditional coloring.

        color_rules: dict mapping (row_idx, col_idx) -> RGBColor
        """
        table = self._add_styled_table(doc, headers, rows)
        if color_rules:
            for (r, c), color in color_rules.items():
                cell = table.rows[r + 1].cells[c]
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = color
                        run.font.bold = True
        return table

    def _fmt(self, val, suffix="", decimals=2):
        if val is None:
            return "N/A"
        if isinstance(val, (int, float)):
            if val == 0 and suffix != "%":
                return "N/A"
            return f"{val:.{decimals}f}{suffix}"
        return str(val)

    def _fmt_cap(self, val):
        if not val or val == 0:
            return "N/A"
        if val > 1e12:
            return f"{val / 1e12:.1f}万亿"
        if val > 1e8:
            return f"{val / 1e8:.1f}亿"
        return f"{val / 1e4:.1f}万"

    def _score_color(self, score):
        """Color for factor score: >70 blue, 40-70 gray, <40 orange."""
        if score >= 70:
            return COLOR_BLUE
        elif score < 40:
            return COLOR_ORANGE
        return None

    def generate_stock_report(
        self,
        symbol: str,
        name: str,
        quote: dict,
        fundamental: dict,
        signals: dict,
        factor_scores: dict[str, float] | None = None,
        ai_commentary: str = "",
        technical_chart_path: str | None = None,
        radar_chart_path: str | None = None,
        valuation_chart_path: str | None = None,
    ) -> str:
        """Generate a complete stock analysis Word document."""
        doc = Document()
        fd = fundamental or {}
        section_num = 0

        # ===== Title =====
        title = doc.add_heading(f"{name}（{symbol}）量化分析报告", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Date subtitle
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(f"报告日期：{datetime.now().strftime('%Y年%m月%d日')}")
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_GRAY

        # Data freshness
        report_date = fd.get("report_date", "")
        if report_date:
            fresh_para = doc.add_paragraph()
            fresh_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                rd = datetime.strptime(str(report_date)[:10], "%Y-%m-%d")
                age_days = (datetime.now() - rd).days
                fresh_text = f"财务数据报告期：{report_date}（距今{age_days}天）"
                run = fresh_para.add_run(fresh_text)
                run.font.size = Pt(9)
                run.font.color.rgb = COLOR_ORANGE if age_days > 180 else COLOR_GRAY
                if age_days > 180:
                    run2 = fresh_para.add_run("  [数据较旧，请关注最新财报]")
                    run2.font.size = Pt(9)
                    run2.font.color.rgb = COLOR_ORANGE
                    run2.font.bold = True
            except (ValueError, TypeError):
                run = fresh_para.add_run(f"财务数据报告期：{report_date}")
                run.font.size = Pt(9)
                run.font.color.rgb = COLOR_GRAY

        # ===== Executive Summary =====
        section_num += 1
        doc.add_heading(f"{'一二三四五六七八'[section_num-1]}、投资摘要", level=1)

        price = quote.get("price", 0)
        change_pct = quote.get("change_pct", 0)
        overall = signals.get("overall", "中性")
        tech_score = signals.get("score", 0)
        avg_factor = (sum(factor_scores.values()) / len(factor_scores)) if factor_scores else 0

        # Build summary text
        change_dir = "上涨" if change_pct > 0 else ("下跌" if change_pct < 0 else "持平")
        trend_text = f"当前价格¥{price:.2f}，今日{change_dir}{abs(change_pct):.2f}%，技术面{overall}（评分{tech_score}）。"

        factor_text = ""
        if factor_scores and avg_factor > 0:
            if avg_factor >= 70:
                factor_text = f"多因子综合得分{avg_factor:.1f}，整体表现优秀。"
            elif avg_factor >= 40:
                factor_text = f"多因子综合得分{avg_factor:.1f}，整体表现中等。"
            else:
                factor_text = f"多因子综合得分{avg_factor:.1f}，整体表现偏弱。"

        risk_flags = []
        pe = quote.get("pe") or fd.get("pe_ttm")
        if pe is not None and pe < 0:
            risk_flags.append("当前处于亏损状态")
        debt = fd.get("debt_ratio")
        if debt is not None and debt > 70:
            risk_flags.append(f"资产负债率较高（{debt:.1f}%）")
        ocfps = fd.get("ocfps")
        eps = fd.get("eps")
        if ocfps is not None and eps is not None and eps > 0 and ocfps < 0:
            risk_flags.append("经营现金流为负")
        risk_text = f"风险关注：{'；'.join(risk_flags)}。" if risk_flags else ""

        summary_para = doc.add_paragraph()
        summary_para.add_run(trend_text).font.size = Pt(11)
        if factor_text:
            summary_para.add_run(factor_text).font.size = Pt(11)
        if risk_text:
            run = summary_para.add_run(risk_text)
            run.font.size = Pt(11)
            run.font.color.rgb = COLOR_ORANGE

        # ===== Realtime Quote =====
        section_num += 1
        doc.add_heading(f"{'一二三四五六七八'[section_num-1]}、实时行情", level=1)

        change_str = f"+{change_pct:.2f}%" if change_pct >= 0 else f"{change_pct:.2f}%"
        change_arrow = "\u2191" if change_pct > 0 else ("\u2193" if change_pct < 0 else "\u2192")
        volume = quote.get("volume", 0)
        volume_str = f"{volume / 10000:.1f}万手" if volume > 10000 else f"{volume:.0f}手"
        amount = quote.get("amount", 0)
        amount_str = f"{amount / 1e8:.2f}亿" if amount > 1e8 else f"{amount / 1e4:.1f}万"

        quote_rows = [
            ["最新价", f"¥{price:.2f}", "涨跌幅", f"{change_arrow} {change_str}"],
            ["今开", self._fmt(quote.get("open")), "最高", self._fmt(quote.get("high"))],
            ["最低", self._fmt(quote.get("low")), "昨收", self._fmt(quote.get("prev_close"))],
            ["成交量", volume_str, "成交额", amount_str],
            ["换手率", self._fmt(quote.get("turnover_rate"), "%"), "总市值", self._fmt_cap(quote.get("total_market_cap") or fd.get("total_mv"))],
        ]
        # Color the change cell
        change_color = COLOR_RED if change_pct > 0 else (COLOR_GREEN if change_pct < 0 else None)
        color_rules = {}
        if change_color:
            color_rules[(0, 3)] = change_color
        self._add_colored_table(doc, ["指标", "数值", "指标", "数值"], quote_rows, color_rules)

        # ===== Financial Metrics =====
        section_num += 1
        doc.add_heading(f"{'一二三四五六七八'[section_num-1]}、核心财务指标", level=1)

        fin_rows = [
            ["PE(TTM)", self._fmt(fd.get("pe_ttm")), "PB(MRQ)", self._fmt(fd.get("pb"))],
            ["PS(TTM)", self._fmt(fd.get("ps_ttm")), "ROE", self._fmt(fd.get("roe"), "%")],
            ["毛利率", self._fmt(fd.get("gross_margin"), "%"), "净利率", self._fmt(fd.get("net_margin"), "%")],
            ["营业利润率", self._fmt(fd.get("operating_margin"), "%"), "营收增长", self._fmt(fd.get("revenue_growth"), "%")],
            ["EPS", self._fmt(fd.get("eps"), "元"), "BPS", self._fmt(fd.get("bps"), "元")],
            ["资产负债率", self._fmt(fd.get("debt_ratio"), "%"), "流动比率", self._fmt(fd.get("current_ratio"))],
        ]
        self._add_styled_table(doc, ["指标", "数值", "指标", "数值"], fin_rows)

        # ===== Earnings Quality =====
        section_num += 1
        doc.add_heading(f"{'一二三四五六七八'[section_num-1]}、现金流与盈利质量", level=1)

        ocf_eps_ratio = None
        quality_label = "N/A"
        quality_color = None
        if ocfps is not None and eps is not None and abs(eps) > 0.001:
            ocf_eps_ratio = ocfps / eps
            if ocf_eps_ratio > 1.0:
                quality_label = "优秀"
                quality_color = COLOR_BLUE
            elif ocf_eps_ratio > 0.5:
                quality_label = "一般"
                quality_color = COLOR_ORANGE
            else:
                quality_label = "存疑"
                quality_color = COLOR_RED
        elif eps is not None and eps < 0:
            quality_label = "亏损"
            quality_color = COLOR_RED

        eq_rows = [
            ["每股经营现金流(OCFPS)", self._fmt(ocfps, "元")],
            ["每股收益(EPS)", self._fmt(eps, "元")],
            ["OCF/EPS比值", self._fmt(ocf_eps_ratio) if ocf_eps_ratio else "N/A"],
            ["盈利质量评估", quality_label],
            ["速动比率", self._fmt(fd.get("quick_ratio"))],
        ]
        eq_color_rules = {}
        if quality_color:
            eq_color_rules[(3, 1)] = quality_color
        self._add_colored_table(doc, ["指标", "数值"], eq_rows, eq_color_rules)

        # Explanation
        eq_note = doc.add_paragraph()
        run = eq_note.add_run(
            "说明：OCF/EPS > 1.0 表示盈利有充足现金流支撑；"
            "0.5~1.0 盈利质量一般；< 0.5 盈利质量存疑，需关注应收账款。"
        )
        run.font.size = Pt(8)
        run.font.color.rgb = COLOR_GRAY
        run.italic = True

        # ===== Historical Valuation Chart =====
        if valuation_chart_path and Path(valuation_chart_path).exists():
            section_num += 1
            doc.add_heading(f"{'一二三四五六七八'[section_num-1]}、历史估值分位", level=1)
            doc.add_picture(valuation_chart_path, width=Inches(6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ===== Technical Analysis =====
        section_num += 1
        doc.add_heading(f"{'一二三四五六七八'[section_num-1]}、技术分析", level=1)

        signal_rows = []
        label_map = {
            "macd_signal": "MACD", "rsi_signal": "RSI", "kdj_signal": "KDJ",
            "boll_signal": "布林带", "ma_alignment": "均线排列",
        }
        for key, label in label_map.items():
            if key in signals:
                signal_rows.append([label, str(signals[key])])
        score = signals.get("score", "N/A")
        overall = signals.get("overall", "N/A")
        signal_rows.append(["综合评分", f"{score}"])
        signal_rows.append(["综合研判", str(overall)])

        # Color the score row
        sig_colors = {}
        score_val = signals.get("score", 0)
        if isinstance(score_val, (int, float)):
            if score_val >= 15:
                sig_colors[(len(signal_rows) - 2, 1)] = COLOR_RED
            elif score_val <= -15:
                sig_colors[(len(signal_rows) - 2, 1)] = COLOR_GREEN
        self._add_colored_table(doc, ["技术指标", "信号"], signal_rows, sig_colors)

        if technical_chart_path and Path(technical_chart_path).exists():
            doc.add_paragraph()
            doc.add_picture(technical_chart_path, width=Inches(6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ===== Factor Scores =====
        if factor_scores:
            section_num += 1
            doc.add_heading(f"{'一二三四五六七八'[section_num-1]}、多因子量化评分", level=1)

            factor_rows = [[k, f"{v:.1f}"] for k, v in factor_scores.items()]
            avg_score = sum(factor_scores.values()) / len(factor_scores)
            factor_rows.append(["综合得分", f"{avg_score:.1f}"])

            # Conditional coloring for factor scores
            fc_rules = {}
            for idx, (_, v) in enumerate(factor_scores.items()):
                c = self._score_color(v)
                if c:
                    fc_rules[(idx, 1)] = c
            avg_c = self._score_color(avg_score)
            if avg_c:
                fc_rules[(len(factor_rows) - 1, 1)] = avg_c

            self._add_colored_table(doc, ["因子维度", "得分(0-100)"], factor_rows, fc_rules)

            if radar_chart_path and Path(radar_chart_path).exists():
                doc.add_paragraph()
                doc.add_picture(radar_chart_path, width=Inches(5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ===== AI Commentary =====
        if ai_commentary:
            section_num += 1
            cn_num = "一二三四五六七八九十"
            num_char = cn_num[section_num - 1] if section_num <= len(cn_num) else str(section_num)
            doc.add_heading(f"{num_char}、AI智能分析研判", level=1)

            for line in ai_commentary.split("\n"):
                line = line.strip()
                if not line:
                    continue
                if line.startswith("##"):
                    doc.add_heading(line.lstrip("#").strip(), level=2)
                elif line.startswith("#"):
                    doc.add_heading(line.lstrip("#").strip(), level=2)
                elif line.startswith("- ") or line.startswith("* "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                elif line.startswith("**") and line.endswith("**"):
                    p = doc.add_paragraph()
                    run = p.add_run(line.strip("*"))
                    run.bold = True
                else:
                    doc.add_paragraph(line)

        # ===== Disclaimer =====
        doc.add_paragraph()
        disclaimer = doc.add_paragraph()
        disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = disclaimer.add_run(
            "免责声明：本报告由AI量化分析系统自动生成，仅供参考，不构成投资建议。"
            "投资有风险，决策需谨慎。"
        )
        run.font.size = Pt(8)
        run.font.color.rgb = COLOR_GRAY
        run.italic = True

        # Save
        filename = f"{name}_{symbol}_量化分析报告.docx"
        filepath = self.output_dir / filename
        doc.save(str(filepath))
        logger.info(f"Report saved to {filepath}")
        return str(filepath)
